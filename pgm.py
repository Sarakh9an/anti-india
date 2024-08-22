import streamlit as st
from eventregistry import EventRegistry, QueryArticlesIter, QueryItems
import json
from docx import Document
from datetime import datetime, timedelta
import io
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Your EventRegistry API key

#secret_key= "030d6822-b0ed-40e0-994e-6b6489a413fd"

# Initialization block
keywords_list = ["India", "USA"]
#, "China", "Bagladesh", "UK", "Lanka", "AFGHANISTAN", "MALDIVES", "MAURITIUS", "TIBET", "IRAN", "ISIS", "AL-QAEDA", "AL-SHABAAB",
 #                 "BOKO", "HARAM", "MODI", "SECURITY", "SIKH", "MILITANT", "CANADA", "AUSTRALIA", "CYBER", "FICN",
  #                 "LAUNDERING", "IB", "R&AW", "South China Sea", "Russia Ukraine", "Indo-Pacific", "BRI", "Nepal", "Taiwan", "Buddhism"]
journal_list = ["washingtonpost.com"]
#, "nytimes.com", "thetimes.co.uk", "ft.com"]
#journal_list = ["deccanherald"]
#["navbharattimes.indiatimes"],"INDIAN EXPRESS", "The Hindu", "Times of India", "sundayguardianlive", "sundaystandard", "hindustantimes", "dailypioneer", "thestatesman", "tribuneindia", "business-standard",
#         "jagran", "jansatta", "thestatesman", "deccanherald", "asianage", "economictimes.indiatime","The Hindu",
 #            "livehindustan", "navbharattimes.indiatimes", "punjabkesari"]
language_list = ["eng", "hin"]
topics = ["Geo Politics", "Military", "Development", "Health", "Business"]
ignore_topic_list = ['hockey', 'cricket', "Bollywood", "Hollywood", "Box office", "Asia Cup", "Games", "sport", "Fashion"]
max_items_per_journal = 5  # Maximum 5 articles per newspaper
max_total_articles = 35  # Maximum total articles to retrieve
days_range = 1
date_end = datetime.now().strftime('%Y-%m-%d')
date_start = (datetime.now() - timedelta(days=days_range)).strftime('%Y-%m-%d')

# Initialize a set to store unique article URLs
unique_article_urls = set()

# Function to determine if an article should be ignored
def should_ignore_article(article):
    title = article.get("title", "").lower()
    body = article.get("body", "").lower()
    ig = ["World Cup", "cricket", "football", "tennis", "match", "Asia cup", "Kohli", "apple", "Asia Cup", "Top News", "badminton"]
    #ig1 = ["badminton", "Grammy", "Asian Games","Cloud Computing", "Top News", "Games", "sport", "Fashion", "Degree", "seafood", "dogs", "Orange", "Housewives", "DevAnand", "Capes", "Coats", "Recommendations", "Advertisement", "Lifestyle", "top-deals", "Amazon Sale", "IPO", "ICC", "Olympics", "Top News", "Box Office", "Live Streaming", "Rohit Shetty", "teaser", "crop", "monsoon", "stock"]
    #ig2 = ["renewable", "crop", "bed bugs", "biofuels" , "yoga", "5G", "See All", "music", "Sufi", "Frog", "FlashFlood", "glacial", "cancer", "Lion", "elephant", "Championships", "Vodka", "mobile app"]
    ignore_keywords = ig
    for keyword in ignore_keywords:
        if keyword.lower() in title or keyword.lower() in body:
            return True
    return False

# Function to retrieve articles for a given keyword and journal
def retrieve_articles(keyword, journal):
    try:
        er = EventRegistry(apiKey=st.secrets["secret_key"], allowUseOfArchive=False)
        q = QueryArticlesIter(
            keywords=QueryItems.OR([keyword]),
            sourceUri=er.getSourceUri(journal),
            lang=QueryItems.OR(language_list),
            ignoreKeywords=QueryItems.OR(ignore_topic_list),
            isDuplicateFilter="skipDuplicates",
            dataType="news",
            dateStart=date_start,
            dateEnd=date_end
        )

        print(f"Number of results for '{keyword}' in '{journal}': {q.count(er)}")

        # Store the articles in the 'articles' list
        articles = []
        count = 0
        for art in q.execQuery(er, sortBy="rel"):
            articles.append(json.dumps(art, indent=4))
            count += 1
            if count >= max_items_per_journal:
                break  # Stop after reaching the maximum number of articles per journal
            if count >= max_total_articles:
                break  # Stop if the maximum total articles limit is reached

        return articles
    except Exception as e:
        print(f"An error occurred for '{keyword}' in '{journal}': {e}")
        return []

def process_keyword_journal_combination(args):
    keyword, journal = args
    articles = retrieve_articles(keyword, journal)
    return articles

def main():
    st.title("Digital Press Clipping Generator")

    # Initialize all_articles as an empty list
    all_articles = []
    doc = Document()

    # Add a button to trigger document generation
    if st.button("Generate Digital Press Clipping"):
        doc, all_articles = generate_document()
        st.success("Document generated successfully!")

    # Display the document content
    if all_articles:
        st.subheader("Generated Document Content:")
        for article in all_articles:
            art = json.loads(article)
            title = art["title"]
            source = art["url"]
            content = art["body"]

            st.markdown(f"**Title:** {title}")
            st.markdown(f"**** {source}")
            st.markdown("****")
            st.markdown(content)
            st.markdown("---")

    # Add a download button for the generated document
    bio = io.BytesIO()
    doc.save(bio)
    if doc:
        st.download_button(
            label="Click here to download the generated document",
            data=bio.getvalue(),
            file_name="DigitalClippings.docx",
            mime="docx"
        )

def generate_document():
    # Fetch and add articles to the document
    doc = Document()  # Initialize the document
    all_articles = []
    total_articles_count = 0

    # Add the front page
    front_page = doc.add_paragraph()
    front_page.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align the text

    title = front_page.add_run("DIGITAL MEDIA PRESS CLIPPING")
    title.bold = True
    title.font.size = Pt(18)

    # Get the current date and day
    current_date = datetime.now().strftime('%Y-%m-%d')
    current_day = datetime.now().strftime('%A')
    print(current_date)
    print(current_day)

    # Add the date and day on the next line
    date_and_day = front_page.add_run(f"\n{current_date} ({current_day})")
    date_and_day.font.size = Pt(14)

    # Add a page break after the front page
    doc.add_page_break()

    # Continue with adding articles
    
    for keyword in keywords_list:
        for journal in journal_list:
            articles = retrieve_articles(keyword, journal)
            for article in articles:
                art = json.loads(article)
                article_url = art.get("url")
                if article_url not in unique_article_urls:
                    all_articles.append(article)
                    unique_article_urls.add(article_url)

            total_articles_count += len(articles)
            if total_articles_count >= max_total_articles:
                break  # Stop if the maximum total articles limit is reached

    for article in all_articles:
        art = json.loads(article)
        title = art["title"]
        source = art["url"]
        content = art["body"]

        # Add article information to the document
        doc.add_heading(title, level=2)
        doc.add_paragraph(f"Source: {source}")
        doc.add_paragraph("Content:").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph(content).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_paragraph("==================")

    return doc, all_articles

if __name__ == "__main__":
    main()
